"""
Multi-Format Document Ingestion Pipeline

This module loads documents from multiple file formats,
extracts textual content, performs chunking, generates
vector embeddings using HuggingFace models, and stores
the resulting vectors in a Qdrant vector database.

Supported Formats:
    - PDF
    - DOCX
    - TXT
    - CSV
    - XLSX / XLS
    - PNG / JPG / JPEG (OCR)
"""

from pathlib import Path
import uuid
import hashlib

import pandas as pd
import pytesseract  # https://docs.langchain.com/oss/python/langchain/rag#returning-source-documents
from PIL import Image
from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_qdrant import QdrantVectorStore

from qdrant_client.models import Distance, VectorParams
from qdrant_client import QdrantClient

# Configuration
DOCUMENTS_DIRECTORY = "data/docs"
COLLECTION_NAME = "multi_documents"
EMBEDDING_MODEL_NAME = "sentence-transformers/all-mpnet-base-v2"


# Document Loaders
def load_pdf_documents(file_path: Path, loaded_documents: list):
    """
    Extract text content from a PDF file.

    Each page is processed individually and stored as a
    separate LangChain Document.
    """
    reader = PdfReader(file_path)

    for page_number, page in enumerate(reader.pages):
        text = page.extract_text() or ""

        if text.strip():
            loaded_documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "file_path": str(file_path),
                        "page": page_number + 1,
                        "file_type": "pdf",
                    },
                )
            )


def load_docx_documents(file_path: Path, loaded_documents: list):
    """
    Extract text from a Microsoft Word document.
    """
    document = DocxDocument(file_path)

    text = "\n".join(
        paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
    )

    loaded_documents.append(
        Document(
            page_content=text,
            metadata={
                "source": file_path.name,
                "file_path": str(file_path),
                "file_type": "docx",
            },
        )
    )


def load_text_documents(file_path: Path, loaded_documents: list):
    """
    Read content from a plain text file.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

        loaded_documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_path.name,
                    "file_path": str(file_path),
                    "file_type": "txt",
                },
            )
        )


def load_csv_documents(file_path: Path, loaded_documents: list):
    """
    Load CSV data and convert it into text format.
    """
    df = pd.read_csv(file_path)

    loaded_documents.append(
        Document(
            page_content=df.to_string(index=False),
            metadata={
                "source": file_path.name,
                "file_path": str(file_path),
                "file_type": "csv",
            },
        )
    )


def load_excel_documents(file_path: Path, loaded_documents: list):
    """
    Load all worksheets from an Excel workbook.
    """
    sheets = pd.read_excel(file_path, sheet_name=None)

    for sheet_name, df in sheets.items():
        loaded_documents.append(
            Document(
                page_content=df.to_string(index=False),
                metadata={
                    "source": file_path.name,
                    "file_path": str(file_path),
                    "sheet": sheet_name,
                    "file_type": "excel",
                },
            )
        )


def load_image_documents(file_path: Path, loaded_documents: list):
    """
    Extract text from images using OCR.
    """
    image = Image.open(file_path)
    extracted_text = pytesseract.image_to_string(image)

    loaded_documents.append(
        Document(
            page_content=extracted_text,
            metadata={
                "source": file_path.name,
                "file_path": str(file_path),
                "file_type": "image",
            },
        )
    )


# Document Loading Orchestrator
def load_all_documents(docs_directory: str):
    """
    Recursively load all supported documents from a directory.
    """
    loaded_documents = []
    documents_path = Path(docs_directory)

    try:
        for current_file in documents_path.rglob("*"):

            if not current_file.is_file():
                continue

            file_extension = current_file.suffix.lower()

            # Handle PDF Documents
            if file_extension == ".pdf":
                load_pdf_documents(current_file, loaded_documents)

            # Handle Docx Documents
            elif file_extension == ".docx":
                load_docx_documents(current_file, loaded_documents)

            # Handle Text Documents
            elif file_extension == ".txt":
                load_text_documents(current_file, loaded_documents)

            # Handle CSV Documents
            elif file_extension == ".csv":
                load_csv_documents(current_file, loaded_documents)

            # Handle Excel Documents
            elif file_extension in [".xlsx", ".xls"]:
                load_excel_documents(current_file, loaded_documents)

            # Handle Image Documents
            elif file_extension in [".png", ".jpg", ".jpeg"]:
                load_image_documents(current_file, loaded_documents)

            else:
                print(f"SKIPPED Unsupported file type:  {current_file.name}")

        print(f"Loaded {len(loaded_documents)} document records.")

    except Exception as error:
        print(f"Failed while processing {current_file.name}: {str(error)}")

    return loaded_documents


# Text Chunking - http://docs.langchain.com/oss/python/integrations/splitters/recursive_text_splitter
def create_document_chunks(documents):
    """
    Split documents into smaller chunks for embedding generation.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=100, chunk_overlap=20, length_function=len, is_separator_regex=False
    )

    document_chunks = text_splitter.split_documents(documents)
    print(f"Generated {len(document_chunks)} document chunks")

    return document_chunks


# Embedding Model
def initialize_embeddings():
    """
    Initialize the HuggingFace embedding model.
    """
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_NAME,
        encode_kwargs={"normalize_embeddings": True},
    )


# Vector Database Storage
def store_documents_in_qdrant(document_chunks, embeddings):
    """
    Store document chunks and embeddings in Qdrant.
    """
    client = QdrantClient(url="http://localhost:6333")

    # Create collection if not exists
    if not client.collection_exists(collection_name=COLLECTION_NAME):
        client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=768, distance=Distance.COSINE),
        )

        print(f"Created collection: {COLLECTION_NAME}")

    # Initialize vector store
    vector_store = QdrantVectorStore(
        client=client, collection_name=COLLECTION_NAME, embedding=embeddings
    )

    document_ids = []

    for chunk in document_chunks:
        unique_content = (
            chunk.metadata.get("source", "")
            + str(chunk.metadata.get("page", ""))
            + str(chunk.metadata.get("sheet", ""))
            + chunk.page_content
        )

        content_hash = hashlib.sha256(
        unique_content.encode()
        ).hexdigest()

        document_ids.append(
            str(uuid.UUID(content_hash[:32]))
        )

    vector_store.add_documents(
        documents=document_chunks,
        ids=document_ids,
    )

    print(
        f"Stored {len(document_chunks)} chunks "
        f"into Qdrant collection {COLLECTION_NAME}"
    )


# Main Ingestion Pipeline
def ingest_documents():
    """
    Execute the complete ingestion workflow.

    Steps:
        1. Load documents
        2. Create chunks
        3. Generate embeddings
        4. Store vectors in Qdrant
    """

    print("STEP 1: Loading documents...")
    loaded_documents = load_all_documents(DOCUMENTS_DIRECTORY)

    print("STEP 2: Creating chunking...")
    document_chunks = create_document_chunks(loaded_documents)

    print("STEP 3: Initializing embeddings...")
    embeddings = initialize_embeddings()

    print("STEP 4: Storing vectors in Qdrant...")
    store_documents_in_qdrant(document_chunks, embeddings)

    print("Document ingestion completed.")


if __name__ == "__main__":
    ingest_documents()
